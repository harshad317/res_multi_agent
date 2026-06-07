from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from research_foundry.models import ResearchReport


DEFAULT_MAIN_CONFERENCE_CHECKS = [
    "Originality against nearest prior work",
    "Significance to the ICLR/NeurIPS community",
    "Technical depth beyond an engineering wrapper",
    "Mechanistic insight or theory of why it works",
    "Strong baseline and SOTA comparison plan",
    "Evaluation breadth across tasks, seeds, and regimes",
    "Statistical reliability and uncertainty reporting",
    "Ablations and negative controls that isolate the contribution",
    "Reproducibility and artifact readiness",
    "Reviewer objection survivability",
    "Main-paper narrative clarity",
    "MVP evidence strong enough to justify scaling",
]


def write_implementation_docx(report: ResearchReport, path: str | Path) -> Path:
    """Create a DOCX implementation plan for the selected idea."""

    output_path = Path(path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_styles(document)

    title = report.selection.selected_title
    document.add_heading(f"Implementation Plan: {title}", level=0)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Research Foundry selected paper direction").italic = True

    _add_metadata_table(document, report)
    _add_selection_summary(document, report)
    _add_execution_brief(document, report)
    _add_markdownish(document, report.implementation_plan.content)
    _add_appendix(document, report)

    document.save(output_path)
    return output_path


def _configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(8)
        code_style.paragraph_format.left_indent = Inches(0.18)
        code_style.paragraph_format.right_indent = Inches(0.1)
        code_style.paragraph_format.space_before = Pt(0)
        code_style.paragraph_format.space_after = Pt(0)


def _add_metadata_table(document: Document, report: ResearchReport) -> None:
    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    values = [
        ("Field", report.request.field),
        ("Objective", report.request.objective),
        ("Target venues", ", ".join(report.request.target_venues)),
        ("Generated", report.generated_at.isoformat()),
    ]
    for idx, (key, value) in enumerate(values):
        table.cell(idx, 0).text = key
        table.cell(idx, 1).text = value


def _add_selection_summary(document: Document, report: ResearchReport) -> None:
    document.add_heading("Selected Idea Decision", level=1)
    document.add_paragraph(report.selection.rationale)
    document.add_heading("Why This Can Be ICLR/NeurIPS-Worthy", level=2)
    document.add_paragraph(report.selection.iclr_neurips_case)
    _add_selected_novelty_audit(document, report)

    if report.selection.decisive_strengths:
        document.add_heading("Decisive Strengths", level=2)
        for item in report.selection.decisive_strengths:
            document.add_paragraph(item, style="List Bullet")

    if report.selection.decisive_risks:
        document.add_heading("Decisive Risks", level=2)
        for item in report.selection.decisive_risks:
            document.add_paragraph(item, style="List Bullet")


def _add_selected_novelty_audit(document: Document, report: ResearchReport) -> None:
    selected_audit = next(
        (
            audit
            for audit in report.novelty_audits
            if audit.idea_title == report.selection.selected_title
        ),
        None,
    )
    if selected_audit is None:
        return

    document.add_heading("Novelty Collision Audit", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Audit Field"
    table.rows[0].cells[1].text = "Value"
    for label, value in [
        ("Main-track verdict", selected_audit.main_track_verdict),
        ("Novelty score", f"{selected_audit.novelty_score}/10"),
        ("Paper-worth score", f"{selected_audit.paper_worth_score}/10"),
        ("Venue-upside score", f"{selected_audit.venue_upside_score}/10"),
        ("Workshop risk", str(selected_audit.workshop_risk)),
    ]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    if selected_audit.closest_prior_work:
        document.add_heading("Closest Prior-Work Collisions", level=3)
        collision_table = document.add_table(rows=1, cols=4)
        collision_table.style = "Table Grid"
        header = collision_table.rows[0].cells
        header[0].text = "Prior Work"
        header[1].text = "Severity"
        header[2].text = "Overlap"
        header[3].text = "Required Differentiator"
        for collision in selected_audit.closest_prior_work:
            row = collision_table.add_row().cells
            row[0].text = collision.prior_work
            row[1].text = collision.severity
            row[2].text = _clean_inline_markdown(collision.overlap)
            row[3].text = _clean_inline_markdown(collision.required_differentiator)

    if selected_audit.main_track_blockers:
        document.add_heading("Main-Track Blockers", level=3)
        for blocker in selected_audit.main_track_blockers:
            document.add_paragraph(_clean_inline_markdown(blocker), style="List Bullet")

    if selected_audit.required_reframing:
        document.add_heading("Required Reframing", level=3)
        for item in selected_audit.required_reframing:
            document.add_paragraph(_clean_inline_markdown(item), style="List Bullet")

    if selected_audit.direct_comparisons_required:
        document.add_heading("Mandatory Direct Comparisons", level=3)
        for item in selected_audit.direct_comparisons_required:
            document.add_paragraph(_clean_inline_markdown(item), style="List Bullet")


def _add_execution_brief(document: Document, report: ResearchReport) -> None:
    document.add_heading("How To Use This Document", level=1)
    document.add_paragraph(
        "This file is a full research blueprint, not a linear coding checklist. Use the "
        "execution brief first, then treat the detailed method, experiments, and appendices "
        "as reference material."
    )

    document.add_heading("Immediate Build Order", level=2)
    next_steps = report.selection.required_next_steps[:6]
    if not next_steps:
        next_steps = [
            "Define the smallest MVP that isolates the core mechanism.",
            "Implement the MVP with strict budget and cost logging.",
            "Run the MVP against the strongest simple baselines.",
            "Inspect failure cases before adding full benchmark scale.",
            "Only start the paper draft after the MVP clears the paper-worthiness gate.",
        ]
    for step in next_steps:
        document.add_paragraph(_clean_inline_markdown(step), style="List Number")

    document.add_heading("Paper-Worthiness Gate", level=2)
    gate_rows = [
        (
            "Research worth continuing",
            "The MVP beats strong simple baselines under matched cost and reveals a real "
            "mechanism, not just prompt luck.",
        ),
        (
            "Paper worth drafting",
            "The full method improves at least one serious SOTA baseline under identical "
            "data, budget, model, and evaluation rules.",
        ),
        (
            "ICLR/NeurIPS target",
            "The method shows a clear, reproducible win across tasks or provides a new "
            "measurement/theory contribution reviewers cannot reduce to prior work.",
        ),
        (
            "Kill or pivot",
            "If the fixed MVP cannot beat random, stratified, and classical racing baselines, "
            "do not scale the experiment until the mechanism is revised.",
        ),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Gate"
    table.rows[0].cells[1].text = "Required Evidence"
    for gate, evidence in gate_rows:
        row = table.add_row().cells
        row[0].text = gate
        row[1].text = evidence

    scores = [
        ("Research worth", report.selection.research_worth_score),
        ("Paper worth", report.selection.paper_worth_score),
        ("Venue upside", report.selection.venue_upside_score),
        ("Fixed-pool only", report.selection.fixed_pool_only_score),
    ]
    if any(value is not None for _, value in scores):
        document.add_heading("Current Upside Scores", level=2)
        score_table = document.add_table(rows=1, cols=2)
        score_table.style = "Table Grid"
        score_table.rows[0].cells[0].text = "Dimension"
        score_table.rows[0].cells[1].text = "Score"
        for label, value in scores:
            row = score_table.add_row().cells
            row[0].text = label
            row[1].text = "Not scored" if value is None else f"{value}/10"

    if report.selection.breakthrough_condition:
        document.add_heading("Breakthrough Condition", level=2)
        document.add_paragraph(_clean_inline_markdown(report.selection.breakthrough_condition))

    _add_main_conference_checklist(document, report)


def _add_main_conference_checklist(document: Document, report: ResearchReport) -> None:
    document.add_heading("ICLR/NeurIPS Main Conference Checklist", level=2)
    document.add_paragraph(
        "Use this grid as the acceptance-bar audit. A serious main-conference target should "
        f"clear {report.request.ambition_floor}/10 on every row, not only on the headline score."
    )
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Check"
    table.rows[0].cells[1].text = "Score"
    table.rows[0].cells[2].text = "Evidence Needed"
    table.rows[0].cells[3].text = "Failure Mode"

    if report.selection.main_conference_checklist:
        for check in report.selection.main_conference_checklist:
            row = table.add_row().cells
            row[0].text = check.name
            row[1].text = f"{check.score}/{check.threshold}"
            row[2].text = _clean_inline_markdown(check.evidence_needed)
            row[3].text = _clean_inline_markdown(check.failure_mode)
        return

    for check_name in DEFAULT_MAIN_CONFERENCE_CHECKS:
        row = table.add_row().cells
        row[0].text = check_name
        row[1].text = "Not scored"
        row[2].text = "Define concrete evidence before treating this idea as submission-grade."
        row[3].text = "Reviewers can reject the paper if this dimension is weak or unsupported."


def _add_markdownish(document: Document, content: str) -> None:
    table_buffer: list[str] = []
    in_code = False

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            _add_markdown_table(document, table_buffer)
            table_buffer = []

    for raw_line in content.splitlines():
        line = raw_line.strip()
        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue

        if in_code:
            document.add_paragraph(raw_line.rstrip() or " ", style="Code Block")
            continue

        if not line:
            flush_table()
            continue

        if _looks_like_table_line(line):
            table_buffer.append(line)
            continue

        flush_table()

        if line.startswith("# "):
            document.add_heading(_clean_inline_markdown(line[2:].strip()), level=1)
        elif line.startswith("## "):
            document.add_heading(_clean_inline_markdown(line[3:].strip()), level=2)
        elif line.startswith("### "):
            document.add_heading(_clean_inline_markdown(line[4:].strip()), level=3)
        elif line.startswith("- "):
            document.add_paragraph(_clean_inline_markdown(line[2:].strip()), style="List Bullet")
        elif _is_numbered_item(line):
            document.add_paragraph(
                _clean_inline_markdown(line.split(".", 1)[1].strip()), style="List Number"
            )
        elif line.startswith(">"):
            document.add_paragraph(_clean_inline_markdown(line.lstrip("> ")))
        else:
            document.add_paragraph(_clean_inline_markdown(line))

    flush_table()


def _add_appendix(document: Document, report: ResearchReport) -> None:
    document.add_page_break()
    document.add_heading("Appendix: Experiment Plan", level=1)
    _add_markdownish(document, report.experiment_plan.content)

    document.add_heading("Appendix: First Reviewer Risks", level=1)
    if report.reviews:
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "Idea"
        header[1].text = "Reviewer"
        header[2].text = "Average"
        header[3].text = "Fatal Flaws"
        for review in report.reviews:
            row = table.add_row().cells
            row[0].text = review.idea_title
            row[1].text = review.reviewer
            row[2].text = str(review.average)
            row[3].text = "; ".join(review.fatal_flaws)

    if report.novelty_audits:
        document.add_heading("Appendix: Novelty Audit For All Ideas", level=1)
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        header = table.rows[0].cells
        header[0].text = "Idea"
        header[1].text = "Verdict"
        header[2].text = "Novelty"
        header[3].text = "Paper"
        header[4].text = "Venue"
        header[5].text = "Closest Collisions"
        for audit in report.novelty_audits:
            row = table.add_row().cells
            row[0].text = audit.idea_title
            row[1].text = audit.main_track_verdict
            row[2].text = f"{audit.novelty_score}/10"
            row[3].text = f"{audit.paper_worth_score}/10"
            row[4].text = f"{audit.venue_upside_score}/10"
            row[5].text = "; ".join(
                f"{collision.prior_work} ({collision.severity})"
                for collision in audit.closest_prior_work
            )

    if report.literature.sources:
        document.add_heading("Appendix: Sources", level=1)
        for source in report.literature.sources:
            text = source.title or source.url
            document.add_paragraph(f"{text}: {source.url}", style="List Bullet")


def _is_numbered_item(line: str) -> bool:
    return bool(re.match(r"^\d+\.\s+\S", line))


def _clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def _looks_like_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def _is_table_separator(line: str) -> bool:
    cells = _split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _split_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return [_clean_inline_markdown(cell) for cell in cells]


def _add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [_split_table_row(line) for line in lines if line.strip()]
    rows = [row for idx, row in enumerate(rows) if idx != 1 or not _is_table_separator(lines[idx])]
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    for col_idx in range(col_count):
        table.rows[0].cells[col_idx].text = rows[0][col_idx] if col_idx < len(rows[0]) else ""

    for data_row in rows[1:]:
        row = table.add_row().cells
        for col_idx in range(col_count):
            row[col_idx].text = data_row[col_idx] if col_idx < len(data_row) else ""
