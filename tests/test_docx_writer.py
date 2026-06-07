from docx import Document

from research_foundry.docx_writer import write_implementation_docx
from research_foundry.models import (
    AgentArtifact,
    IdeaCandidate,
    ResearchReport,
    ResearchRequest,
    SelectionDecision,
)


def _minimal_report() -> ResearchReport:
    request = ResearchRequest(
        field="prompt optimization",
        objective="find breakthrough ideas",
        dry_run=True,
    )
    idea = IdeaCandidate(
        title="Test Idea",
        thesis="A thesis.",
        core_mechanism="A mechanism.",
        novelty_claim="A novelty claim.",
        why_reviewers_care="A significance claim.",
    )
    selection = SelectionDecision(
        selected_title="Test Idea",
        rationale="Best available.",
        iclr_neurips_case="Strong if validated.",
        research_worth_score=8,
        paper_worth_score=8,
        venue_upside_score=8,
        fixed_pool_only_score=8,
        breakthrough_condition="The MVP must beat the strongest simple baseline.",
        score=8,
    )
    implementation = AgentArtifact(
        agent_name="Implementation Architect",
        model="test",
        content="""# Method

```python
# 1. Warm start.
print("x")
```

| Col A | Col B |
| --- | --- |
| A | B |
""",
    )
    artifact = AgentArtifact(agent_name="x", model="test", content="x")
    return ResearchReport(
        request=request,
        literature=artifact,
        novelty_gaps=[],
        ideas=[idea],
        reviews=[],
        selection=selection,
        experiment_plan=artifact,
        implementation_plan=implementation,
        final_recommendation=artifact,
    )


def test_docx_writer_preserves_fenced_code_blocks(tmp_path):
    output = tmp_path / "implementation.docx"
    write_implementation_docx(_minimal_report(), output)

    doc = Document(output)
    headings = {
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.style.name.startswith("Heading")
    }
    code_lines = [
        paragraph.text
        for paragraph in doc.paragraphs
        if paragraph.style.name == "Code Block"
    ]

    assert "# 1. Warm start." in code_lines
    assert "# 1. Warm start." not in headings
    assert any(paragraph.text == "How To Use This Document" for paragraph in doc.paragraphs)
    assert len(doc.tables) >= 2
